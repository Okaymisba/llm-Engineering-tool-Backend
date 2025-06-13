import io
from docx import Document
import os

def parse_docx(file_input):
    """
    Parses the content of a .docx file and extracts its text content.

    Args:
        file_input: Either a file path (str) or a file stream (bytes/io.BytesIO)
        
    Returns:
        str: The extracted text from the DOCX file
    """
    try:
        # Handle both file paths and file streams
        if isinstance(file_input, str):
            # If it's a file path, open the file
            document = Document(file_input)
        else:
            # If it's a stream, use it directly
            document = Document(io.BytesIO(file_input))

        text = ""
        for paragraph in document.paragraphs:
            text += paragraph.text + "\n"

        return text.strip()
        
    except Exception as e:
        raise Exception(f"Error processing DOCX: {str(e)}")